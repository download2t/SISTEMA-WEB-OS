from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Contrato
from .forms import ContratoForm
from datetime import timedelta, timezone
from django.contrib.auth.decorators import login_required, user_passes_test
from core.views import has_permission
from django.contrib.auth.models import Group
from django.shortcuts import render
from django.utils.timezone import now
from django.shortcuts import render
from datetime import timedelta
from .models import Contrato
from django.contrib.auth.models import Group,  User 
from django.http import HttpResponse
from reportlab.platypus import Table, TableStyle # type: ignore
from django.contrib import messages
from reportlab.pdfgen import canvas # type: ignore
from django.utils import timezone
from reportlab.lib import colors # type: ignore
from reportlab.lib.pagesizes import landscape, letter # type: ignore
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from reportlab.lib.units import inch # type: ignore
from docx import Document # type: ignore
import openpyxl # type: ignore
from openpyxl.utils import get_column_letter # type: ignore
from django.http import HttpResponse
from django.db.models import Q
from django.utils.timezone import now
from django.conf import settings # Importar settings

@login_required
def listar_contratos(request):
    contratos = Contrato.objects.all()

    # Obtendo parâmetros da requisição
    tipo_data = request.GET.get("tipo_data", "assinatura")
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")
    razao_social = request.GET.get("razao_social", "").strip()
    usuario_responsavel_id = request.GET.get("usuario_responsavel") 
    status = request.GET.get("status")

    # Filtro por Data
    if data_inicio and data_fim:
        if tipo_data == "assinatura":
            contratos = contratos.filter(data_assinatura__range=[data_inicio, data_fim])
        elif tipo_data == "validade":
            contratos = contratos.filter(data_validade__range=[data_inicio, data_fim])

    # Filtro por Razão Social, Nome Fantasia ou Documento (busca parcial)
    if razao_social:
        contratos = contratos.filter(
            Q(nome_fantasia__icontains=razao_social) |
            Q(razao_social__icontains=razao_social) |
            Q(documento__icontains=razao_social)
        )

    # Filtro por Usuário Responsável
    if usuario_responsavel_id: # Usar o novo nome da variável
        contratos = contratos.filter(usuario_responsavel=usuario_responsavel_id) # Filtrar pelo novo campo

    # Filtro por Status (ativo, inativo, vencido ou vencendo)
    hoje = now().date()
    if status == "ativo":
        contratos = contratos.filter(ativo=True)
    elif status == "inativo":
        contratos = contratos.filter(ativo=False)
    elif status == "vencido":
        contratos = contratos.filter(data_validade__lt=hoje)
    elif status == "vencendo_60_dias":
        data_limite = hoje + timedelta(days=60)
        contratos = contratos.filter(data_validade__range=[hoje, data_limite], ativo=True)

    # Pegando todos os usuários para exibir no filtro
    # Ordenar por username para um dropdown mais organizado
    usuarios = User.objects.all().order_by('username')

    context = {
        "contratos": contratos,
        "usuarios": usuarios, # Alterado de 'grupos' para 'usuarios'
        "tipo_data": tipo_data,
        "data_inicio": data_inicio,
        "data_fim": data_fim,
        "razao_social": razao_social,
        "usuario_selecionado": usuario_responsavel_id, # Passa o ID do usuário selecionado
        "status": status,
    }

    return render(request, "contratos/listar_contratos.html", context)


def has_permission_contratos(user):
    # Verifica se o usuário está autenticado e é staff ou pertence aos grupos "ADMIN" ou "LIDERANÇA - Mudar"
    return user.is_authenticated and (user.is_staff or user.groups.filter(name__in=['ADMINISTRATIVO', 'ALMOXARIFADO','ADMIN','COMPRAS',
                                                            'COMERCIAL','GOVERNANÇA','LIDERANÇA','TI','RESERVAS','RECEPÇÃO']).exists())
# contratos/views.py

@login_required
@user_passes_test(has_permission_contratos, login_url='403')
def criar_contrato(request):
    """ Cria um novo contrato, garantindo preenchimento correto """
    if request.method == 'POST':
        form = ContratoForm(request.POST, request.FILES) 
        if form.is_valid():
            form.save()
            messages.success(request, "Contrato criado com sucesso!")
            return redirect('listar_contratos')
        else:
            # Adicione este print para depuração de erros do formulário
            print("Erros do formulário ao criar:", form.errors)
            messages.error(request, "Houve um erro ao criar o contrato. Verifique os campos.")
    else:
        form = ContratoForm()
    
    try:
        max_upload_mb = int(settings.FILE_UPLOAD_MAX_MEMORY_SIZE / (1024 * 1024))
    except (TypeError, ZeroDivisionError):
        max_upload_mb = 20 # Valor padrão caso haja problema na configuração

    context = {
        'form': form,
        'titulo': 'Criar Contrato',
        'max_upload_mb': max_upload_mb
    }
    return render(request, 'contratos/criar_contrato.html', context)

@login_required
@user_passes_test(has_permission_contratos, login_url='403')
def editar_contrato(request, contrato_id):
    """ Edita um contrato existente """
    contrato = get_object_or_404(Contrato, id=contrato_id)

    if request.method == 'POST':
        form = ContratoForm(request.POST, request.FILES, instance=contrato)
        if form.is_valid():
            form.save()
            messages.success(request, "Contrato atualizado com sucesso!")
            return redirect('listar_contratos')
        else:
            # Adicione este print para depuração de erros do formulário
            print("Erros do formulário ao editar:", form.errors)
            messages.error(request, "Houve um erro ao atualizar o contrato. Verifique os campos.")
    else:
        form = ContratoForm(instance=contrato)
    
    try:
        max_upload_mb = int(settings.FILE_UPLOAD_MAX_MEMORY_SIZE / (1024 * 1024))
    except (TypeError, ZeroDivisionError):
        max_upload_mb = 20 # Valor padrão caso haja problema na configuração

    context = {
        'form': form,
        'titulo': 'Editar Contrato',
        'max_upload_mb': max_upload_mb
    }
    return render(request, 'contratos/editar_contrato.html', context)


@login_required
def visualizar_contrato(request, contrato_id):
    """ Visualiza os detalhes de um contrato existente. """
    contrato = get_object_or_404(Contrato, id=contrato_id)

    context = {
        'contrato': contrato,
        'titulo': 'Detalhes do Contrato'
    }
    return render(request, 'contratos/visualizar_contrato.html', context)

@login_required
@user_passes_test(has_permission, login_url='403')  # Garantindo que o usuário tenha permissão
def inativar_contrato(request, contrato_id):
    """ Inativa um contrato existente (não exclui o contrato, apenas marca como inativo) """
    contrato = get_object_or_404(Contrato, id=contrato_id)
    
    # Inativa o contrato
    contrato.ativo = False
    contrato.save()

    # Mensagem de sucesso
    messages.success(request, "Contrato inativado com sucesso!")

    return redirect('listar_contratos')

def ativar_contrato(request, contrato_id):
    """ Ativa um contrato existente (não exclui o contrato, apenas marca como ativo) """
    contrato = get_object_or_404(Contrato, id=contrato_id)
    
    # Ativa o contrato
    contrato.ativo = True
    contrato.save()

    # Mensagem de sucesso
    messages.success(request, "Contrato ativado com sucesso!")

    return redirect('listar_contratos')

def contratos_vencendo(request):
    hoje = now().date()
    limite_vencimento = hoje + timedelta(days=30)
    contratos = Contrato.objects.filter(ativo=True, data_validade__lte=limite_vencimento, data_validade__gte=hoje)
    
    return render(request, 'contratos/listar_contratos.html', {
        'contratos': contratos,
        'titulo_pagina': 'Contratos Vencendo'
    })


from django.views.decorators.http import require_POST

@login_required
@require_POST # Ensures this view only accepts POST requests
@user_passes_test(has_permission_contratos, login_url='403')
def toggle_status_contrato(request, contrato_id):
    """ Toggles the 'ativo' status of a contract. """
    contrato = get_object_or_404(Contrato, id=contrato_id)
    
    # Flip the boolean status
    contrato.ativo = not contrato.ativo
    contrato.save()

    # Create a success message
    if contrato.ativo:
        messages.success(request, f"O contrato '{contrato.nome_fantasia or contrato.razao_social}' foi ATIVADO com sucesso.")
    else:
        messages.warning(request, f"O contrato '{contrato.nome_fantasia or contrato.razao_social}' foi DESATIVADO.")
        
    # Redirect back to the edit page
    return redirect('editar_contrato', contrato_id=contrato.id)
####################################### RELATÓRIOS #######################################


def abreviar_texto(texto, max_length):
    if texto is None:
        return ""
    if not isinstance(texto, str):
        texto = str(texto)
    return texto if len(texto) <= max_length else texto[:max_length - 3] + "..."

# --- Views Django ---
@login_required
@user_passes_test(has_permission, login_url='403')
def listar_contratos_rel(request):
    contratos = Contrato.objects.all()

    razao_social = request.GET.get('razao_social')
    nome_fantasia = request.GET.get('nome_fantasia')
    documento = request.GET.get('documento')
    usuario_responsavel_id = request.GET.get('usuario_responsavel')
    data_inicio = request.GET.get('data_inicio')
    data_fim = request.GET.get('data_fim')
    status = request.GET.get('status')
    valor_min = request.GET.get('valor_min')
    valor_max = request.GET.get('valor_max')

    if razao_social:
        contratos = contratos.filter(razao_social__icontains=razao_social)
    if nome_fantasia:
        contratos = contratos.filter(nome_fantasia__icontains=nome_fantasia)
    if documento:
        contratos = contratos.filter(documento__icontains=documento)
    if usuario_responsavel_id:
        contratos = contratos.filter(usuario_responsavel_id=usuario_responsavel_id)
    if data_inicio:
        contratos = contratos.filter(data_assinatura__gte=data_inicio)
    if data_fim:
        contratos = contratos.filter(data_assinatura__lte=data_fim)
    if status:
        if status == 'ativo':
            contratos = contratos.filter(ativo=True)
        elif status == 'inativo':
            contratos = contratos.filter(ativo=False)
    if valor_min:
        try:
            valor_min = float(valor_min.replace(',', '.'))
            contratos = contratos.filter(valor__gte=valor_min)
        except ValueError:
            pass
    if valor_max:
        try:
            valor_max = float(valor_max.replace(',', '.'))
            contratos = contratos.filter(valor__lte=valor_max)
        except ValueError:
            pass

    # Esta linha já está correta, buscando todos os usuários e ordenando-os
    usuarios = User.objects.all().order_by('first_name', 'last_name', 'username')

    return render(request, 'contratos/relatorios.html', {
        'contratos': contratos,
        'usuarios': usuarios,  # Esta linha já está passando os usuários para o template
    })


def gerar_relatorio_pdf(request):
    contratos = Contrato.objects.all()

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="relatorio_contratos.pdf"'
    pdf = canvas.Canvas(response, pagesize=landscape(letter))

    largura_pagina, altura_pagina = landscape(letter)
    margem_horizontal = 30
    margem_vertical_topo = 70
    margem_vertical_base = 20

    linha_inicial_conteudo = altura_pagina - margem_vertical_topo
    linha_final_conteudo = margem_vertical_base + 20

    def desenhar_cabecalho(pagina_atual, total_paginas):
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(margem_horizontal, altura_pagina - 30, "Relatório de Contratos")
        pdf.setFont("Helvetica", 10)
        if request.user.is_authenticated:
            usuario_info = f"{request.user.get_full_name() or request.user.username} (ID: {request.user.id})"
        else:
            usuario_info = "Usuário Desconhecido"
        pdf.drawString(margem_horizontal, altura_pagina - 50, f"Gerado por: {usuario_info}")
        pdf.drawRightString(largura_pagina - margem_horizontal, altura_pagina - 30, f"Página {pagina_atual} de {total_paginas}")

    def desenhar_rodape():
        data_geracao = now().strftime('%d/%m/%Y %H:%M:%S')
        pdf.setFont("Helvetica", 9)
        pdf.drawString(margem_horizontal, 30, f"Data/Hora: {data_geracao}")

    def desenhar_tabela(dados, y_pos):
        table = Table(dados, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e5e7eb')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
            ('BACKGROUND', (0, 2), (-1, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#4b5563')),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
        ]))
        table_width = largura_pagina - 2 * margem_horizontal
        table.wrapOn(pdf, table_width, altura_pagina)
        table.drawOn(pdf, margem_horizontal, y_pos)

    largura_tabela = largura_pagina - 2 * margem_horizontal
    col_widths = [
        largura_tabela * 0.18,  # Nome Fantasia
        largura_tabela * 0.10,  # Documento
        largura_tabela * 0.10,  # Telefone
        largura_tabela * 0.20,  # E-mail (Manter 0.20, mas abreviar mais)
        largura_tabela * 0.08,  # Data de Validade
        largura_tabela * 0.20,  # Usuário Responsável
        largura_tabela * 0.08,  # Valor
        largura_tabela * 0.06,  # Status
    ]

    cabecalho_tabela = [["Nome Fantasia", "Documento", "Telefone", "E-mail", "Validade", "Usuário Responsável", "Valor", "Status"]]
    data = []
    for contrato in contratos:
        data.append([
            abreviar_texto(contrato.nome_fantasia or contrato.razao_social, 25),
            contrato.documento,
            contrato.telefone or '',
            # AQUI: Diminuí o max_length para E-mail de 35 para 25 (ou 20 se ainda for muito)
            abreviar_texto(contrato.email or '', 25), # Ajustado para 25 caracteres
            contrato.data_validade.strftime("%d/%m/%Y"),
            abreviar_texto(contrato.usuario_responsavel.get_full_name() or contrato.usuario_responsavel.username, 30),
            f'R$ {contrato.valor:.2f}',
            'ATIVO' if contrato.ativo else 'INATIVO',
        ])

    itens_por_pagina = 20
    paginas = [data[i:i + itens_por_pagina] for i in range(0, len(data), itens_por_pagina)]

    total_paginas = len(paginas)
    for i, pagina in enumerate(paginas):
        desenhar_cabecalho(i + 1, total_paginas)

        pagina_com_cabecalho = cabecalho_tabela + pagina

        temp_table = Table(pagina_com_cabecalho, colWidths=col_widths)
        temp_table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ]))

        table_width = largura_pagina - 2 * margem_horizontal
        table_height = temp_table.wrapOn(pdf, table_width, 0)[1]

        y_pos = linha_inicial_conteudo - table_height

        if y_pos < linha_final_conteudo and i < len(paginas) - 1:
            pdf.showPage()
            desenhar_cabecalho(i + 1, total_paginas)
            y_pos = linha_inicial_conteudo - table_height

        desenhar_tabela(pagina_com_cabecalho, y_pos)
        desenhar_rodape()

        if i < len(paginas) - 1:
            pdf.showPage()

    pdf.save()
    return response


def generate_word(request):
    contratos = Contrato.objects.all()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="relatorio_contratos.docx"'

    document = Document()
    document.add_heading('Relatório de Contratos', 0)

    for contrato in contratos:
        document.add_paragraph(f"Razão Social: {contrato.razao_social} | Nome Fantasia: {contrato.nome_fantasia or ''} | Documento: {contrato.documento}")
        document.add_paragraph(f"Telefone: {contrato.telefone or ''} | E-mail: {contrato.email or ''}")
        document.add_paragraph(f"Data de Assinatura: {contrato.data_assinatura.strftime('%d/%m/%Y')} | Data de Validade: {contrato.data_validade.strftime('%d/%m/%Y')}")
        document.add_paragraph(f"Usuário Responsável: {contrato.usuario_responsavel.get_full_name() or contrato.usuario_responsavel.username} | Valor: R$ {contrato.valor:.2f} | Ativo: {'Sim' if contrato.ativo else 'Não'}")
        document.add_paragraph(f"Descrição: {contrato.descricao or ''}")
        document.add_paragraph("-" * 100)

    document.save(response)
    return response


def generate_excel(request):
    contratos = Contrato.objects.all()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="relatorio_contratos.xlsx"'

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Contratos"

    headers = ['Razão Social', 'Nome Fantasia', 'Documento', 'Telefone', 'E-mail', 'Assinatura', 'Validade', 'Usuário Responsável', 'Valor', 'Ativo', 'Descrição']
    ws.append(headers)

    for contrato in contratos:
        ws.append([
            contrato.razao_social,
            contrato.nome_fantasia or '',
            contrato.documento,
            contrato.telefone or '',
            contrato.email or '',
            contrato.data_assinatura.strftime('%d/%m/%Y'),
            contrato.data_validade.strftime('%d/%m/%Y'),
            contrato.usuario_responsavel.get_full_name() or contrato.usuario_responsavel.username,
            float(contrato.valor),
            'Sim' if contrato.ativo else 'Não',
            contrato.descricao or '',
        ])

    for col_num, column_title in enumerate(headers, 1):
        column_letter = get_column_letter(col_num)
        ws.column_dimensions[column_letter].auto_size = True

    wb.save(response)
    return response
